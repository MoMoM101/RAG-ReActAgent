from rag.loaders import LOADERS, load_document, load_docx, load_image, load_pdf, load_xlsx


class TestLoadDocx:
    def test_basic(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        text = load_docx(str(f))
        assert "第一段内容" in text
        assert "第二段内容" in text

    def test_empty(self, tmp_path):
        from docx import Document
        doc = Document()
        f = tmp_path / "empty.docx"
        doc.save(str(f))

        text = load_docx(str(f))
        assert text == ""


class TestLoadXlsx:
    def test_basic(self, tmp_path):
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.append(["姓名", "年龄"])
        ws.append(["张三", 25])
        ws.append(["李四", 30])
        f = tmp_path / "test.xlsx"
        wb.save(str(f))

        text = load_xlsx(str(f))
        assert "张三" in text
        assert "李四" in text
        assert "姓名" in text

    def test_empty(self, tmp_path):
        from openpyxl import Workbook

        wb = Workbook()
        f = tmp_path / "empty.xlsx"
        wb.save(str(f))

        assert load_xlsx(str(f)) == ""

    def test_escapes_markdown_cells(self, tmp_path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["名称", "说明"])
        ws.append(["项目|甲", "第一行\n第二行"])
        f = tmp_path / "escaped.xlsx"
        wb.save(str(f))

        text = load_xlsx(str(f))

        assert r"项目\|甲" in text
        assert "第一行<br>第二行" in text


class TestLoadPdf:
    def test_text_pdf(self, tmp_path):
        from config import settings
        old = settings.ocr_enabled
        settings.ocr_enabled = False
        try:
            import fitz
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Hello World from PDF", fontname="helv", fontsize=12)
            f = tmp_path / "test.pdf"
            doc.save(str(f))
            doc.close()

            text = load_pdf(str(f))
            assert "Hello World from PDF" in text
        finally:
            settings.ocr_enabled = old


class TestLoadImage:
    def test_ocr_disabled(self, tmp_path):
        from config import settings
        old = settings.ocr_enabled
        settings.ocr_enabled = False
        try:
            from PIL import Image
            img = Image.new("RGB", (100, 100), color="white")
            f = tmp_path / "test.png"
            img.save(str(f))

            text = load_image(str(f))
            assert text == ""
        finally:
            settings.ocr_enabled = old


class TestRouting:
    def test_all_formats_have_loader(self):
        for ext in [".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".jpg", ".jpeg", ".png"]:
            assert ext in LOADERS, f"Missing loader for {ext}"

    def test_load_document_routes_correctly(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("routing test", encoding="utf-8")
        text = load_document(str(f), ".txt")
        assert text == "routing test"

        f2 = tmp_path / "test.md"
        f2.write_text("# MD routing", encoding="utf-8")
        text = load_document(str(f2), ".md")
        assert text == "# MD routing"
